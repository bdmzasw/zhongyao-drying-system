from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

import streamlit as st
import pandas as pd
import sys
import os
from datetime import datetime
import base64
from io import BytesIO
import matplotlib
import matplotlib.pyplot as plt

# ========== 全局修复中文乱码（线上永久生效） ==========
plt.rcParams['axes.unicode_minus'] = False
matplotlib.rcParams['font.sans-serif'] = ['WenQuanYi Zen Hei', 'SimHei', 'DejaVu Sans']
matplotlib.rcParams['axes.unicode_minus'] = False

sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

# ===================== 安全转换 =====================
def safe_float(value, default=0.0):
    if pd.isna(value) or value == "-":
        return default
    s = str(value).strip()
    if "~" in s or "-" in s:
        try:
            sep = "~" if "~" in s else "-"
            a, b = s.split(sep)
            return (float(a.strip()) + float(b.strip())) / 2
        except:
            return default
    try:
        return float(s)
    except:
        return default

# ===================== 读取CSV =====================
def read_csv_safe(path):
    try:
        return pd.read_csv(path, encoding="utf-8-sig")
    except:
        return pd.read_csv(path, encoding="gbk")

@st.cache_data
def load_all_data():
    return {
        "herbs": read_csv_safe("data/药材库.csv"),
        "techs": read_csv_safe("data/技术库.csv"),
        "regions": read_csv_safe("data/区域库.csv"),
        "carbon": read_csv_safe("data/区域碳排放因子.csv"),
        "lcc": read_csv_safe("data/LCC经济成本库-干燥设备对比.csv"),
        "kinetics": read_csv_safe("data/干燥动力学.csv")
    }

data = load_all_data()
herbs_df = data["herbs"]
regions_df = data["regions"]
techs_df = data["techs"]

# ===================== 页面 =====================
st.set_page_config(page_title="报告导出", page_icon="📄", layout="wide")
st.markdown("""
<style>
[data-testid="stSidebarNav"] {display: none;}
.stApp { background-color: #f5f9f5; }
</style>
""", unsafe_allow_html=True)

if st.button("🏠 返回主页"):
    st.switch_page("Home.py")

# ===================== 侧边栏 ======================
with st.sidebar:
    st.subheader("⚙️ 总操作仪表盘")
    herb_list = ["请选择"] + herbs_df["药材标准名称(药典名)"].dropna().unique().tolist()
    selected_herb = st.selectbox("🌿 药材品种", herb_list, index=herb_list.index(st.session_state.get("selected_herb", "请选择")) if st.session_state.get("selected_herb") in herb_list else 0)
    region_list = ["请选择"] + regions_df["产区名称"].dropna().unique().tolist()
    selected_area = st.selectbox("📍 产区", region_list, index=region_list.index(st.session_state.get("selected_area", "请选择")) if st.session_state.get("selected_area") in region_list else 0)
    electricity_price = st.number_input("⚡ 电价（元/kWh）", value=st.session_state.get("electricity_price", 0.6), step=0.01)
    annual_capacity = st.number_input("📦 年处理量（吨/年）", value=st.session_state.get("annual_capacity", 400), step=50)

    st.session_state["selected_herb"] = selected_herb
    st.session_state["selected_area"] = selected_area
    st.session_state["electricity_price"] = electricity_price
    st.session_state["annual_capacity"] = annual_capacity
    st.markdown("---")
    st.caption("✅ 全系统同步生效")

# ===================== 主界面 =====================
st.markdown("## 📄 中药材干燥工艺决策报告")
st.caption("实时预览与 Word 导出完全一致（含文字 + 图表 + 双方案）")
st.divider()

if selected_herb == "请选择" or selected_area == "请选择":
    st.warning("⚠️ 请先选择药材与产区再生成报告")
else:
    herb_info = herbs_df[herbs_df["药材标准名称(药典名)"] == selected_herb].iloc[0]
    region_info = regions_df[regions_df["产区名称"] == selected_area].iloc[0]
    province = region_info["所辖主要省市"].split("、")[0]
    init_mc = safe_float(herb_info["鲜品初始含水率(%)"])
    final_mc = safe_float(herb_info["药典规定成品含水率(%)"])
    water_removed = 1000 * (init_mc - final_mc) / (100 - final_mc)

    # 计算所有工艺
    result = []
    for _, t in techs_df.iterrows():
        en = safe_float(t["单位能耗(kWh/kg水)"]) * water_removed
        cost = en * electricity_price
        carbon = en * 0.55
        result.append({
            "干燥技术": t["干燥技术"],
            "成分保留率(%)": round(safe_float(t["有效成分保留率(%)"]),1),
            "干燥时间(h)": round(safe_float(t["干燥时间范围 (h)"]),1),
            "能耗(kWh/吨)": round(en,1),
            "碳排放(kgCO₂/吨)": round(carbon,1),
            "加工成本(元/吨)": round(cost,1),
        })

    df = pd.DataFrame(result)
    df["综合得分"] = (df["成分保留率(%)"]*0.4 - df["能耗(kWh/吨)"]*0.2 - df["干燥时间(h)"]*0.15 - df["碳排放(kgCO₂/吨)"]*0.15 - df["加工成本(元/吨)"]*0.1).round(2)
    df = df.sort_values("综合得分", ascending=False)
    best1 = df.iloc[0]
    best2 = df.iloc[1]

    # 生成图片并保存到内存
    fig1, ax1 = plt.subplots(figsize=(6, 3.5))
    ax1.barh(df["干燥技术"], df["成分保留率(%)"], color="#4a9f75")
    ax1.set_xlabel("成分保留率 (%)")
    ax1.set_title("各干燥工艺有效成分保留率对比")
    plt.tight_layout()
    buf1 = BytesIO()
    fig1.savefig(buf1, dpi=150, format="png")
    buf1.seek(0)

    fig2, ax2 = plt.subplots(figsize=(6, 3.5))
    ax2.bar(df["干燥技术"], df["综合得分"], color="#3b82f6")
    ax2.set_ylabel("综合得分")
    ax2.set_title("各工艺综合得分对比")
    plt.xticks(rotation=20, ha="right")
    plt.tight_layout()
    buf2 = BytesIO()
    fig2.savefig(buf2, dpi=150, format="png")
    buf2.seek(0)

    # ===================== 报告预览 =====================
    with st.expander("📄 报告实时预览", expanded=True):
        st.markdown(f"# 中药材干燥工艺决策报告")
        st.markdown(f"**生成时间**：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        st.markdown(f"**药材**：{selected_herb}　|　**产地**：{selected_area}　|　**年处理量**：{annual_capacity} 吨")
        st.divider()

        st.markdown("## 一、药材基础信息")
        st.write(f"- 初始含水率：{init_mc:.1f} %")
        st.write(f"- 药典成品含水率：{final_mc:.1f} %")
        st.write(f"- 每吨脱水量：{water_removed:.1f} kg")

        st.divider()
        st.markdown("## 二、双方案推荐")
        c1, c2 = st.columns(2)
        with c1:
            st.success("🏆 主方案（综合最优）")
            st.write(f"工艺：{best1['干燥技术']}")
            st.write(f"得分：{best1['综合得分']}")
            st.write(f"保留率：{best1['成分保留率(%)']}%")
        with c2:
            st.info("📌 备选方案（次优）")
            st.write(f"工艺：{best2['干燥技术']}")
            st.write(f"得分：{best2['综合得分']}")
            st.write(f"保留率：{best2['成分保留率(%)']}%")

        st.divider()
        st.markdown("## 三、工艺对比图表")
        col1, col2 = st.columns(2)
        with col1:
            st.pyplot(fig1)
        with col2:
            st.pyplot(fig2)

        st.divider()
        st.markdown("## 四、年度能耗与碳排放")
        st.write(f"- 年耗电量：{best1['能耗(kWh/吨)'] * annual_capacity:.0f} kWh")
        st.write(f"- 年碳排放：{best1['碳排放(kgCO₂/吨)'] * annual_capacity / 1000:.2f} tCO₂")
        st.write(f"- 年能源成本：{best1['加工成本(元/吨)'] * annual_capacity:.0f} 元")

        st.divider()
        st.markdown("## 五、结论与建议")
        st.write(f"1. 推荐主方案：{best1['干燥技术']}")
        st.write(f"2. 设备受限时可选：{best2['干燥技术']}")
        st.write("3. 工艺兼顾品质、能耗、效率与低碳要求。")

    # ===================== Word 导出（含图片 + 全部内容）=====================
    def generate_full_docx():
        doc = Document()

        # ========== 修复 Word 中文乱码（核心代码）==========
        style = doc.styles['Normal']
        style.font.name = 'SimSun'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        style.font.size = Pt(12)

        doc.add_heading("中药材干燥工艺决策报告", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"药材：{selected_herb}　|　产地：{selected_area}　|　年处理量：{annual_capacity}吨")

        doc.add_heading("一、药材基础信息", level=1)
        doc.add_paragraph(f"初始含水率：{init_mc:.1f}%")
        doc.add_paragraph(f"药典要求含水率：{final_mc:.1f}%")
        doc.add_paragraph(f"每吨脱水量：{water_removed:.1f}kg")

        doc.add_heading("二、双方案推荐", level=1)
        doc.add_heading("主方案", level=2)
        doc.add_paragraph(f"工艺：{best1['干燥技术']}  得分：{best1['综合得分']}")
        doc.add_paragraph(f"保留率：{best1['成分保留率(%)']}%  能耗：{best1['能耗(kWh/吨)']}kWh/吨")

        doc.add_heading("备选方案", level=2)
        doc.add_paragraph(f"工艺：{best2['干燥技术']}  得分：{best2['综合得分']}")
        doc.add_paragraph(f"保留率：{best2['成分保留率(%)']}%  能耗：{best2['能耗(kWh/吨)']}kWh/吨")

        doc.add_heading("三、工艺对比图表", level=1)
        doc.add_picture(buf1, width=Inches(5.0))
        doc.add_paragraph("图1 各工艺有效成分保留率对比").alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_picture(buf2, width=Inches(5.0))
        doc.add_paragraph("图2 各工艺综合得分对比").alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_heading("四、年度能耗与碳排放", level=1)
        doc.add_paragraph(f"年耗电量：{best1['能耗(kWh/吨)'] * annual_capacity:.0f} kWh")
        doc.add_paragraph(f"年碳排放：{best1['碳排放(kgCO₂/吨)'] * annual_capacity / 1000:.2f} tCO₂")
        doc.add_paragraph(f"年能源成本：{best1['加工成本(元/吨)'] * annual_capacity:.0f} 元")

        doc.add_heading("五、结论与建议", level=1)
        doc.add_paragraph(f"优先选用：{best1['干燥技术']}")
        doc.add_paragraph(f"备选方案：{best2['干燥技术']}")
        doc.add_paragraph("本方案可实现品质、能耗、效率、低碳协同优化。")

        final_buf = BytesIO()
        doc.save(final_buf)
        final_buf.seek(0)
        return final_buf

    # 下载按钮
    docx_file = generate_full_docx()
    b64 = base64.b64encode(docx_file.getvalue()).decode()
    href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="中药材干燥报告_{selected_herb}.docx">📥 下载完整Word报告（含文字+图表）</a>'
    st.markdown(href, unsafe_allow_html=True)
    st.success("✅ Word 报告已包含：全文 + 两张图表 + 双方案")

st.divider()
